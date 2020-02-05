from guizero import App, Text, PushButton, CheckBox, TextBox, ButtonGroup, Picture
import subprocess
import datetime
import cv2
import numpy as np
import pandas as pd
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
import math
import os
def addressess():
    #shows 5 icons, address message, background image and start trip button
    #hides all possible widgets from previous screens
    show_all(),hide_all(),address_message.show(), start_trip.show(), logo_picture.show()
    #opens flight controller, Q Ground Control is not in GUI subsection
    subprocess.Popen(['C:\Program Files (x86)\QGroundControl\QGroundControl.exe'])
def addressestwo():
    #hides start trip button and shows finish triip button instead
    start_trip.hide(), finish_trip.show()
    #grabs the time the trip started
    set_beg_time()
def adduser():
    #checks if the new username is valid
    user_name_used = repeated_user()
    #opens username txt file
    user_list = open("user.txt", "a")
    #opens password textfile
    pass_list = open("pass.txt", "a")
    #if username is valid enter if statement
    if user_name_used == True:
        #hide error message
        create_user_error.hide()
        #add user name to file
        user_list.write(user_input_box.value)
        #add an "enter space" to file
        user_list.write("\n")
        #add password to file
        pass_list.write(pass_input_box.value)
        #add "enter space" to file
        pass_list.write("\n")
        #finds name of new file about to write to
        firstlast=user_input_box.value + ".txt"
        #creates this file
        user_info = open(firstlast, "a")
        #adds their username
        user_info.write(user_input_box.value)
        # add "enter space" to file
        user_info.write("\n")
        #adds their password
        user_info.write(pass_input_box.value)
        # add "enter space" to file
        user_info.write("\n")
        #adds their first name
        user_info.write(firstname_input_box.value)
        # add "enter space" to file
        user_info.write("\n")
        #adds their last name
        user_info.write(lastname_input_box.value)
        # add "enter space" to file
        user_info.write("\n")
        #add their location to qground control
        user_info.write(qground_input_box.value)
        # add "enter space" to file
        user_info.write("\n")
        #close document
        user_info.close()
        #creates name of txt file for stats
        stattext = user_input_box.value + "_stat.txt"
        #creates txt file for stats
        stat_text_info = open(stattext, "a")
        #0 as first stats
        stat_text_info.write(str(0))
        #seperate stats with newline
        stat_text_info.write("\n")
        # 0 as first stats
        stat_text_info.write(str(0))
        # seperate stats with newline
        stat_text_info.write("\n")
        # 0 as first stats
        stat_text_info.write(str(0))
        # seperate stats with newline
        stat_text_info.write("\n")
        # 0 as first stats
        stat_text_info.write(str(0))
        # seperate stats with newline
        stat_text_info.write("\n")
        # 0 as first stats
        stat_text_info.write(str(0))
        # seperate stats with newline
        stat_text_info.write("\n")
        #close file
        stat_text_info.close()
        #if username was not valid go to else block
    else:
        #shows error message and allows user to try again
        create_user_error.show()
def create_user():
    #hides previous screen
    signin_message.hide(), username_input_box.hide(), username_message.hide(), password_input_box.hide(), password_message.hide(), spacing_signin1.hide(), signin_button.hide(), createuser.hide()
    #show 5 messages of options when creating account
    create_user_message.show(), firstname_message.show(),  lastname_message.show(), user_message_create.show(), password_message_create.show(),  qground_location_message.show()
    #shows the text box where user actually types too,
    firstname_input_box.show(), lastname_input_box.show(), user_input_box.show(), pass_input_box.show(), qground_input_box.show()
    #changes color to match app's theme
    create_user_message.bg= "slategray4"
    #shows button that will cause code to call add_user
    createuser_finished.show()
def docss():
    #hides all previous buttons and 5 icons
    hide_all(),settings_button.hide(),manual_button.hide(),address_button.hide(),doc_button.hide(),profile_button.hide()
    #shows message alerting user they are on documents
    documents_message.show()
    #gets user's txt file
    a = username_input_box.value + ".txt"
    #creates a vector thatuser txt file will be stored on
    user_info_ary = []
    #opens user text file
    with open(a) as f:
        #calls to begin a for loop to start the process of putting txt file on array
        for line in f:
            #appends vector line by line
            user_info_ary.append(line)
    #closes user txt file
    f.close()
    #declares length to be the size of user stat file
    length = len(user_info_ary)
    #if lengthi s 6 only one doc to show
    if(length == 6):
        #show first doc
        trip0.show()
        #vertical spacing between doc and icons
        trip_space0.show()
        #text of button is changed to last item in doc
        trip0.text = user_info_ary[length-1].strip()
        #adjust height to fit spacing
        trip_space0.height = 214
        #if text file = 7, 2 docs
    elif(length == 7):
        #show 2 buttons and spacing
        trip0.show(), trip1.show(), trip_space0.show()
        #top doc is recent
        trip0.text = user_info_ary[length - 1].strip()
        #next doc is next recent
        trip1.text = user_info_ary[length - 2].strip()
        #changes document spacing
        trip_space0.height = 187
        #if length is 8 there are three documents
    elif(length == 8):
        #show buttons and spacing
        trip0.show(), trip1.show(), trip2.show(), trip_space0.show()
        #change text to be recent document
        trip0.text = user_info_ary[length - 1].strip()
        #change text to be next most recent documemnt
        trip1.text = user_info_ary[length - 2].strip()
        #change text to be the oldest doc
        trip2.text = user_info_ary[length - 3].strip()
        #change height to have spacing fit well
        trip_space0.height = 161
        #if length =9, there are 4 documents
    elif(length == 9):
        #show buttons for documents
        trip0.show(), trip1.show(), trip2.show(), trip3.show(),trip_space0.show()
        #trip0 is most recent document
        #trip1 is next most recent document
        #trip2 is the next oldest
        #trip3is oldest
        trip0.text = user_info_ary[length - 1].strip()
        trip1.text = user_info_ary[length - 2].strip()
        trip2.text = user_info_ary[length - 3].strip()
        trip3.text = user_info_ary[length - 4].strip()
        #changest spacing
        trip_space0.height = 135
        #if length is 10 there are 5 docs
    elif (length == 10):
        #shows buttons for documents
        trip0.show(), trip1.show(), trip2.show(), trip3.show(), trip4.show(),trip_space0.show()
        # trip0 is most recent document
        # trip1 is next most recent document
        # trip2 is the next oldest
        #trip3 is the next oldest
        # trip4 is oldest
        trip0.text = user_info_ary[length - 1].strip()
        trip1.text = user_info_ary[length - 2].strip()
        trip2.text = user_info_ary[length - 3].strip()
        trip3.text = user_info_ary[length - 4].strip()
        trip4.text = user_info_ary[length - 5].strip()
        #changes spacing
        trip_space0.height = 109
        #if length is 11 there are 6 docs to show
    elif (length == 11):
        #shows 6 documents and videos
        trip0.show(), trip1.show(), trip2.show(), trip3.show(), trip4.show(), trip5.show(), trip_space0.show()
        # trip0 is most recent document
        # trip1 is next most recent document
        # trip2 is the next oldest
        # trip3 is the next oldest
        #trip4 is next oldest
        # trip5 is oldest
        trip0.text = user_info_ary[length - 1].strip()
        trip1.text = user_info_ary[length - 2].strip()
        trip2.text = user_info_ary[length - 3].strip()
        trip3.text = user_info_ary[length - 4].strip()
        trip4.text = user_info_ary[length - 5].strip()
        trip5.text = user_info_ary[length - 6].strip()
        #changes spacing
        trip_space0.height = 83
        #if length is 12, i will show 7 options for docs and vids
    elif (length == 12):
        #show options for 7 docs and videos
        trip0.show(), trip1.show(), trip2.show(), trip3.show(), trip4.show(), trip5.show(),trip6.show(), trip_space0.show()
        # trip0 is most recent document
        # trip1 is next most recent document
        # trip2 is the next oldest
        # trip3 is the next oldest
        # trip4 is the next oldest
        #trip5 is the next oldest
        #trip7 is oldest
        trip0.text = user_info_ary[length - 1].strip()
        trip1.text = user_info_ary[length - 2].strip()
        trip2.text = user_info_ary[length - 3].strip()
        trip3.text = user_info_ary[length - 4].strip()
        trip4.text = user_info_ary[length - 5].strip()
        trip5.text = user_info_ary[length - 6].strip()
        trip6.text = user_info_ary[length - 7].strip()
        #changes spacing
        trip_space0.height = 57
        #if length is above 12 show the information on last 8 trips
    elif (length > 12):
        #show button options
        trip0.show(), trip1.show(), trip2.show(), trip3.show(), trip4.show(), trip5.show(),trip6.show(),trip7.show(),trip_space0.show()
        # trip0 is most recent document
        # trip1 is next most recent document
        # trip2 is the next oldest
        # trip3 is the next oldest
        # trip4 is the next oldest
        # trip5 is the next oldest
        # trip7 is the next oldest
        #trip8 is the next oldest
        trip0.text = user_info_ary[length - 1].strip()
        trip1.text = user_info_ary[length - 2].strip()
        trip2.text = user_info_ary[length - 3].strip()
        trip3.text = user_info_ary[length - 4].strip()
        trip4.text = user_info_ary[length - 5].strip()
        trip5.text = user_info_ary[length - 6].strip()
        trip6.text = user_info_ary[length - 7].strip()
        trip7.text = user_info_ary[length - 8].strip()
        #change spcing
        trip_space0.height = 28
        #else there will be no docs to show
    else:
        #change height to fit whole screen
        trip_space0.height = 240
        #show spacing text
        trip_space0.show()
    #show icons
    settings_button1.show(),manual_button1.show(),address_button1.show(),doc_button1.show(),profile_button1.show()
def hide_all():
    #hides widgtes that could have been on previous screen
    settings_message.hide(), username_input_box.hide(), password_input_box.hide(), username_message.hide()
    # hides widgtes that could have been on previous screen
    password_message.hide(),password_message.hide(),createuser.hide(),signin_button.hide(), signin_message.hide()
    # hides widgtes that could have been on previous screen
    signin_message.hide(),spacing_signin1.hide(),video_pic_message.hide(),video_pic_button.hide(),whole_security_message.hide()
    # hides widgtes that could have been on previous screen
    whole_security_button.hide(),manual_message.hide(),address_message.hide(),start_trip.hide(),logo_picture.hide()
    # hides widgtes that could have been on previous screen
    finish_trip.hide(),space_set.hide(),documents_message.hide(),trip0.hide(),trip1.hide(),trip2.hide(),trip3.hide()
    # hides widgtes that could have been on previous screen
    trip4.hide(),trip5.hide(),trip6.hide(),trip7.hide(),settings_button1.hide(),manual_button1.hide(),address_button1.hide()
    # hides widgtes that could have been on previous screen
    doc_button1.hide(), profile_button1.hide(),man0.hide(),mans1.hide(),man2.hide(), man3.hide(),man4.hide(),man5.hide()
    # hides widgtes that could have been on previous screen
    man6.hide(),man7.hide(),man8.hide(),man9.hide(),man10.hide(),man11.hide(),packages.hide(),packages1.hide()
    # hides widgtes that could have been on previous screen
    packages2.hide(),packages3.hide(),packages4.hide(),trip_space0.hide(),profile_message.hide(),trip_total_message.hide()
    # hides widgtes that could have been on previous screen
    distance_total_message.hide(),speed_total_message.hide(),threats_total_message.hide(),logout_button.hide()
    # hides widgtes that could have been on previous screen
    space_profile.hide(), man1.hide(), mannext.hide(), man3.hide(), mans2.hide(), time_total_message.hide()
    #if profile view was changed to be one we need to hide the following widgets
    if profile_view == 1:
        #hides items from the profile
        num_trip.hide(),num_time.hide(),num_speed.hide(), num_distance.hide(), num_threat.hide()
def log_out():
    #shows a black screen when logged out
    #user must restart program to reuse program
    hide_all()
def make_doc():
    #gets doc name
    set_doc_name()
    #declares doc_date to be used globally
    global doc_date
    #creates a document that could be microsoft word
    document = Document()
    #creates text of heading of post trip report
    doc_heading = 'Ground Reconnaissance Robot' + "\n" + 'Contributors: Cody, Ryan, Pulakit, Yousef' + "\n" + "Texas A&M University"
    #creates heading #1(0)
    document.add_heading(doc_heading, 0)
    #blanks doc_date
    doc_date = ""
    #adds a paragraoh
    p = document.add_paragraph(doc_date)
    #make textsaya date
    doc_date = "Date: "
    #make text bold
    p.add_run(doc_date).bold = True
    #update doc_date with date
    doc_date = str(datetime.datetime.now().date()) + "\n"
    #adds it to report
    p.add_run(doc_date)
    #adds "Time in bold to report
    p.add_run("Time: ").bold=True
    #doc date = first half of date
    doc_date= str(datetime.datetime.now().hour) + ":" + str(datetime.datetime.now().minute) + "\n"
    #add the time
    p.add_run(doc_date)
    #add distance in bold
    p.add_run("Distance: ").bold = True
    #shows meters with 0 decimal points
    doc_date =str(format(distance, '.0f')) + " meters" + "\n"
    #adds the meters
    p.add_run(doc_date)
    #adds time taken in bolf
    p.add_run("Time Taken: ").bold = True
    #doc_date = timem taken
    doc_date =str(time_difference) + " seconds" + "\n"
    #adds time to document
    p.add_run(doc_date)
    #adds stats to documents
    p.add_run("Average Speed: ").bold = True
    # adds stats to documents
    doc_date = str(math.trunc(distance / time_difference)) + " m/s" + "\n"
    # adds stats to documents
    p.add_run(doc_date)
    # adds stats to documents
    p.add_run("Desired Number of Waypoints: ").bold = True
    # adds stats to documents
    doc_date = str(len(long_plot_data) - 1) + "\n"
    # adds stats to documents
    p.add_run(doc_date)
    # adds stats to documents
    p.add_run("Starting Point: (").bold = True
    # adds stats to documents
    doc_date = str(long_first_data[0]) + " , " + str(lat_first_data[0]) + " )" + "\n"
    # adds stats to documents
    p.add_run(doc_date)
    # adds stats to documents
    p.add_run("Desired Waypoints: ").bold = True
    # adds stats to documents
    doc_date = "\n"
    # adds stats to documents
    p.add_run(doc_date)
    #creates an iterator
    i = 1
    #begins a while loop to make sections in reports
    while i < len(long_first_data):
        #creates a heading
        doc_date =  "( " + str(long_first_data[i]) + " , " + str(lat_first_data[i]) + " )" + "\n"
        #creates a heading
        p.add_run(doc_date)
        #iterates
        i = i + 1
    # adds stats to documents
    p.add_run("Potential Hazard Locations: ").bold = True
    # adds stats to documents
    p.add_run("\n")
    #iterator
    i=1
    #iterator
    j=1
    #creates empty vector
    haz_long = []
    #creates empty vector
    haz_lat = []
    #looking for hazards
    while i<len(long_first_data):
        #best case of no hazards
        if((long_first_data[i] == long_end_data[j]) & (lat_first_data[i] == lat_end_data[i])):
            #i dont need to do anything in this block so i pass through it
            #was simpler to break if else into this and not just one if
            pass
        else:
            #iterates through end data until hazard was avoided
            while ((long_end_data[j] != long_first_data[i]) | (lat_end_data[j] != lat_first_data[i])):
                #appends hazard to hazard vector
                haz_long.append(long_end_data[j])
                # appends hazard to hazard vector
                haz_lat.append(lat_end_data[j])
                #iterates
                j = j + 1
        # iterates
        i = i + 1
        # iterates
        j = j + 1
    #iterator
    i = 0
    #adds potential hazard points to post trip report
    while i<len(haz_long):
        #formats potential hazards in ( long , lat ) format
        doc_date = "( " + str(haz_long[i]) + " , " + str(haz_lat[i]) + " )" + "\n"
        #adds potential hazards to document
        p.add_run(doc_date)
        #iterates through while loop
        i = i + 1
    #makes the reports style normal
    style = document.styles['Normal']
    #declares a variable that will be used to change font and font size
    font = style.font
    #makes post travel report arial
    font.name = 'Arial'
    #makes post travel font size 10 in word
    font.size = Pt(10)
    #makes the report normal style
    p.style = document.styles['Normal']
    #adds a header that will show map of desired waypoints in that section
    document.add_heading("Desired Trip")
    #adds map with just desired points
    document.add_picture('map_with_points.png', width=Inches(4))
    #adds a page break to document
    #helps with spacing
    document.add_page_break()
    #heading for the photo of map of the actual trip
    document.add_heading("Actual Trip")
    #map with points taken
    document.add_picture('map_with_points_two.png', width=Inches(4))
    #heading for map with desired vs actual route overlayed onto each other
    document.add_heading("Desired vs Actual Trip")
    #shows map of actual vs desired wyapoints
    document.add_picture("map_three.png", width=Inches(4))
    #adds a page break to report for spacing reasons
    document.add_page_break()
    #if the robot goes from more than a to b
    #so if it goes from a to b to c
    #additional information iwll be posted
    if(len(long_first_data) > 2):
        #iterator
        i = 1
        #iterator
        j = 1
        #going to make each a to b to c's information on document
        while i <len(long_first_data):
            #creates empty vector here
            #very important placement, creates an empty vector each iteration
            long_plot = []
            lat_plot = []
            #creates the name of each sub travel's map
            pic_name = "map" + str(i) + ".png"
            #cretes the string for the headings
            head_wp = "From Wavepoint " + str(i - 1) + " to Wavepoint " + str(i)
            #adds each sub travels header
            document.add_heading(head_wp, level=1)
            #creates a new paragraph each iteration
            p = document.add_paragraph("")
            #distance traveled follows the pythagorean theorem
            dist = (((long_first_data[i] - long_first_data[i - 1]) ** 2) + ((lat_first_data[i] - lat_first_data[i - 1]) ** 2)) ** 0.5
            #converts distance into meters
            dist = dist * 111139
            #THIS LINE CREATES FAKE DATA
            #I created this to be able to test code
            #in 404 I will recieve information similar to this
            time_wp = time_difference / (len(long_first_data) - 1)
            #adds distance to paragraph
            p.add_run("Distance: ").bold = True
            #truncates the distance and adds units to string
            doc_date = str(math.trunc(dist)) + " meters" + "\n"
            #addds the actual distance
            p.add_run(doc_date)
            #adds "TIme taken: " in bold on the dooc
            p.add_run("Time taken: ").bold = True
            #creates string with time taken and units
            doc_date = str(time_wp) + " seconds" + "\n"
            #adds string of data to document
            p.add_run(doc_date)
            #average speed in bolf
            p.add_run("Average Speed: ").bold = True
            #creates string containing data
            doc_date = str(math.trunc(dist / time_wp)) + " m/s" + "\n"
            #adds data to document
            p.add_run(doc_date)
            #creates map for no avoidance
            if (lat_first_data[i] == lat_first_data[j])&(long_first_data[i]==long_end_data[j]):
                #dx and dy are total distance on map for long and lat
                #i subtract the min of each from the data
                #finds a percent of how far on the map it should be
                #then later will plot this
                long_plot.append((long_first_data[i-1] - x2) / dx)
                long_plot.append((long_first_data[i] - x2)/dx)
                lat_plot.append((lat_first_data[i-1] - y1) / dy)
                lat_plot.append((lat_first_data[i] - y1) / dy)
                #close the previous plot
                plt.close()
                #makes scatterplot with the data from above
                plt.scatter(long_plot, lat_plot, zorder=1)
                #connects data with line,
                #next semester there will be a lot more data
                #but fr now it works with the fake data given
                plt.plot(long_plot, lat_plot)
                #loads up map
                img = plt.imread("map.png")
                #adjust the values on the side of map
                ext = [0, 1, 0, 1]
                #adds values to map
                plt.imshow(img, zorder=0, extent=ext)
                #makes it precise
                aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
                #final step in making map
                plt.gca().set_aspect(aspect)
                #saves image under pic name
                plt.savefig(pic_name)
                #adds to document there were no obstaclels detected
                p.add_run("No hazard detected.").bold = True
                #adds picture to document
                document.add_picture(pic_name, width=Inches(4))
                #transitions to else block this is for an obstacle was detected
            else:
                #finds percent of data on map in longitude direction
                long_plot.append((long_first_data[i-1]-x2)/dx)
                #finds percent of data  on map in latitude direction
                lat_plot.append((lat_end_data[i-1] - y1)/dy)
                #while still not at designated location
                while ((long_end_data[j] != long_first_data[i]) | (lat_end_data[j] != lat_first_data[i])):
                    # finds percent of data on map in longitude direction
                    long_plot.append((long_end_data[j]-x2)/dx)
                    # finds percent of data  on map in latitude direction
                    lat_plot.append((lat_end_data[j]-y1)/dy)
                    #iterate through actual data
                    j = j + 1
                # finds percent of data  on map in longitude direction
                long_plot.append((long_first_data[i] - x2) / dx)
                # finds percent of data  on map in latitude direction
                lat_plot.append((lat_first_data[i] - y1) / dy)
                #closes plot
                plt.close()
                #scatters point plot
                plt.scatter(long_plot, lat_plot, zorder=1)
                #connects points
                plt.plot(long_plot, lat_plot)
                #opens college station map
                img = plt.imread("map.png")
                #sets boundaries
                ext = [0, 1, 0, 1]
                #shows boundaries
                plt.imshow(img, zorder=0, extent=ext)
                #makes aspect
                aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
                #sets aspect
                plt.gca().set_aspect(aspect)
                #saves map
                plt.savefig(pic_name)
                #made if else statement to have correct grammar
                # if length = 3 one hazard detected
                if(len(lat_plot) == 3):
                    #add hazard warning
                    p.add_run("One Potential Hazard Detected.").bold = True
                    #else hazard is plural
                else:
                    #adds number of hazards
                    p.add_run(str(len(lat_plot) - 2)).bold = True
                    #adds hazard message
                    p.add_run(" Potential Hazards Detected.").bold = True
                    #adds picture
                document.add_picture(pic_name, width=Inches(4))
                #iterates
            i = i + 1
            #iterates
            j = j + 1
        #ADD paragraph here
    #saves document
    document.save(doc_name)
    #makes string of users file
    a = username_input_box.value + ".txt"
    #opens user file
    g = open(a, "a+")
    #makes string of docname
    doc_n = doc_name + "\n"
    #adds docname to user profile
    g.write(doc_n)
    #closes file
    g.close()
    h = username_input_box.value + "_stat.txt"
    update_stat = []
    with open(h) as k:
        #iterates through file
        for line in k:
            #adds stats to vector
            update_stat.append(line)
        length = len(update_stat)
        new_tripu = str(1 + int(update_stat[length - 5].strip())) + "\n"
        timeu = str(time_difference + int(update_stat[length-4].strip())) + "\n"
        distanceu = str(int(float(distance + float(update_stat[length - 3].strip())))) + "\n"
        speedu = int(float(distanceu.strip()) / int(timeu.strip()))
        speedu = str(speedu) + "\n"
        hazu = str(int(update_stat[length - 1].strip()) + len(lat_end_data) - len(lat_first_data)) + "\n"
        k.close()
        write_to_stat = open(h, "a")
        write_to_stat.write(new_tripu)
        write_to_stat.write(timeu)
        write_to_stat.write(distanceu)
        write_to_stat.write(speedu)
        write_to_stat.write(hazu)
        write_to_stat.close()
def make_map():
    #iterator
    i = 0
    #length is used in while loopp
    length = len(lat_first_data)
    #fill vectors with plotting data of DESIRED waypoints
    while i < length:
        #finds vector for longitude plot data; min = 0 max = 1
        long_plot_data.append((long_first_data[i] - x2) / dx)
        # finds vector for latitude plot data; min = 0 max = 1
        lat_plot_data.append((lat_first_data[i] - y1) / dy)
        #iterates to next element
        i = i + 1
    #places dots on map
    plt.scatter(long_plot_data, lat_plot_data, zorder=1)
    #connects data on map
    plt.plot(long_plot_data, lat_plot_data)
    #loads map
    img = plt.imread("map.png")
    #sets boundaries of maps
    ext = [0, 1, 0, 1]
    #places load map on plot
    plt.imshow(img, zorder=0, extent=ext)
    #adjust aspect of plot
    aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
    #Last step of showing map
    plt.gca().set_aspect(aspect)
    #saved map with desired points
    plt.savefig("map_with_points.png", bbox_inches="tight")
    #closes plot so i can do next pllot
    plt.close()
    #i is an iterator
    i = 0
    #length of actual wavepoints gone
    length = len(lat_end_data)
    #sets empty array for actual wavepoints
    lat_plot_two = []
    #sets empty array for latitude wavepoints
    long_plot_two = []
    #will get data for actual wavepoints
    while i < length:
        # finds vector for longitude plot data; min = 0 max = 1
        long_plot_two.append((long_end_data[i] - x2) / dx)
        # finds vector for latitude plot data; min = 0 max = 1
        lat_plot_two.append((lat_end_data[i] - y1) / dy)
        #iterates to next element
        i = i + 1
    #scatter points from actual trip
    plt.scatter(long_plot_two, lat_plot_two, zorder=1)
    #connects dots from actual trip
    plt.plot(long_plot_two, lat_plot_two)
    #loads image
    img = plt.imread("map.png")
    #sets boundaries of map
    ext = [0, 1, 0, 1]
    #shows loaded image
    plt.imshow(img, zorder=0, extent=ext)
    #set aspect
    aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
    #last step of map
    plt.gca().set_aspect(aspect)
    #saves actual wavepoint map
    plt.savefig("map_with_points_two.png", bbox_inches="tight")
    #close plot for future plots
    plt.close()
    #places actual data on red on the plot that will have actual vs desired plot
    plt.scatter(long_plot_two, lat_plot_two, zorder=1, c="red")
    #connects actual data
    plt.plot(long_plot_two, lat_plot_two, c="red")
    #loads image
    img = plt.imread("map.png")
    #sets boundaries
    ext = [0, 1, 0, 1]
    #shows loaded image
    plt.imshow(img, zorder=0, extent=ext)
    #defines aspect of map
    aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
    #sets aspet
    plt.gca().set_aspect(aspect)
    #scatters desired wavepoints in blue
    plt.scatter(long_plot_data, lat_plot_data, zorder=1, c="blue" )
    #connects desired wavepoints in blue
    plt.plot(long_plot_data, lat_plot_data, c="blue")
    #loads image
    img = plt.imread("map.png")
    #sets boundaries
    ext = [0, 1, 0, 1]
    #show sloaded image
    plt.imshow(img, zorder=0, extent=ext)
    #defines aspect
    aspect = img.shape[0] / float(img.shape[1]) * ((ext[1] - ext[0]) / (ext[3] - ext[2]))
    #sets aspect
    plt.gca().set_aspect(aspect)
    #save actual vs desired wavepoint
    plt.savefig("map_three.png")
def profilee():
    #global num_threat, num_speed, num_distance, num_trip, num_time
    hide_all(), profile_button.hide(), doc_button.hide(), address_button.hide(), settings_button.hide(), manual_button.hide()
    #shows second set of icons that are lower than regular
    settings_button1.show(), manual_button1.show(), address_button1.show(), doc_button1.show(), profile_button1.show()
    #shows messages for stats
    profile_message.show(), trip_total_message.show(), time_total_message.show()
    #shows messages for stats
    distance_total_message.show(), speed_total_message.show(), threats_total_message.show(), logout_button.show(), space_profile.show()
    #adjust spacing of profile
    space_profile.height = 152
    #makes user stat file a string
    stat_file = username_input_box.value + "_stat.txt"
    #creates empty vector for stats to be saved in
    user_s_ary = []
    #opens profile stats file
    with open(stat_file) as g:
        #iterates through file
        for line in g:
            #adds stats to vector
            user_s_ary.append(line)
    #closes file
    g.close()
    #defines length of stats array
    length = len(user_s_ary)
    #creates global text so that hide_all() has access to close them in other screens
    global num_trip, num_time, num_distance,num_speed, num_threat
    #number of trips is the fifrth to last thing stored in array
    num_trip = Text(app,text=user_s_ary[length-5].strip(), grid=[4, 2], size=10, align="left")
    #time stat is fourth to last stat in array
    num_time = Text(app, text=user_s_ary[length-4].strip(), grid=[4, 4], size=10, align="left")
    #distance is thrid to last element in array
    num_distance = Text(app, text=user_s_ary[length-3].strip(), grid=[4, 8], size=10, align="left")
    #speed stat is second to last element
    num_speed = Text(app, text=user_s_ary[length-2].strip(), grid=[4, 16], size=10, align="left")
    #threat stat is last thing stored in vector
    num_threat = Text(app, text= user_s_ary[length-1].strip(), grid=[4, 32], size=10, align="left")
    #shows text of all stats
    num_trip.show(), num_time.show(), num_distance.show(), num_speed.show(), num_threat.show()
    #creates profile view boolean to tell if profile has been seen, since none of the above stat text messages
    #have been initiated in beginning, a function that calls hide_all() would result in error, that is why
    #i added a if profile_view ==1 then hide them
    global profile_view
    #lets code know profile has been viewed
    profile_view = 1
def set_beg_time():
    #declares that the following vairables are gloabal
    #this will change the data all functions see the variables
    global time_begin_year, time_begin_month, time_begin_day, time_begin_hour, time_begin_minute, time_begin_second, time_beg
    #sets the year when the program starts
    time_begin_year = datetime.datetime.now().year
    #sets the month when the program starts
    time_begin_month = datetime.datetime.now().month
    # sets the day when the program starts
    time_begin_day = datetime.datetime.now().day
    # sets the hour when the program starts
    time_begin_hour = datetime.datetime.now().hour
    # sets the minute when the program starts
    time_begin_minute = datetime.datetime.now().minute
    # sets the second when the program starts
    time_begin_second = datetime.datetime.now().second
    #finds the seconds that have passed through the day, will be used to find out total time during trip
    time_beg = 60 * 60 * time_begin_hour + 60 * time_begin_minute + time_begin_second
def set_distance():
    #allows every function to see distance changed
    global distance
    #iterator
    i = 1
    #will be used to flag when to stop while loop
    length = len(long_end_data)
    #will find distance traveled in trip
    while i < length:
        #finds a
        dx1 = (long_end_data[i] - long_end_data[i - 1])
        #finds b
        dy1 = (lat_end_data[i] - lat_end_data[i - 1])
        #c = sqrt(a^2 + b^2)
        distance = distance + (dx1 ** 2 + dy1 ** 2) ** 0.5
        #iterates to next element
        i = i + 1
    # converts to meters from degrees
    distance = distance * 111139
def set_doc_name():
    #lets all function see changes to the name of the document
    global doc_name
    #name is "year_month_day_time_hour_minute_second.docx
    doc_name = str(datetime.datetime.now().year) + "_" + str(datetime.datetime.now().month) + "_" + str(
        datetime.datetime.now().day) + "_time_" + str(datetime.datetime.now().hour) + "_" + str(
        datetime.datetime.now().minute) + "_" + str(datetime.datetime.now().second) + ".docx"
def set_end_time():
    # declares that the following vairables are gloabal
    # this will change the data all functions see the variables
    global time_end_year, time_end_month, time_end_day, time_end_hour, time_end_minute,time_end_month, time_end
    # sets the year when the program ends
    time_end_year = datetime.datetime.now().year
    # sets the month when the program ends
    time_end_month = datetime.datetime.now().month
    # sets the day when the program ends
    time_end_day = datetime.datetime.now().day
    # sets the hour when the program ends
    time_end_hour = datetime.datetime.now().hour
    # sets the minute when the program ends
    time_end_minute = datetime.datetime.now().minute
    # sets the second when the program ends
    time_end_second = datetime.datetime.now().second
    #converts to seconds that have occured throughout the day at end of trip
    #will be used to find time used in trip
    time_end = 60 * 60 * time_end_hour + 60 * time_end_minute + time_end_second
def show_all():
    #shows the main 5 bottom icon that for most part are always on
    settings_button.show(), doc_button.show(), address_button.show(), manual_button.show(), profile_button.show()
def set_time_difference():
    #declares time difference to be global
    global time_difference
    #uses equation to find time difference
    time_difference = (time_end - time_beg)
    #takes care of exception
    #started program at 11:59 PM end program at 12:01 AM
    if time_difference < 0:
        #finds time remaining in day and then add to the time finished the next day
        time_difference = 24*60*60 - time_beg + time_end
def trips():
    #trips is all about  making doucments
    #calls function to fnd distance traveled
    set_distance()
    #makes map for desired, actual, and desired vs actual
    make_map()
    #sets end time is called
    set_end_time()
    #finds time total
    set_time_difference()
    #makes the document
    make_doc()
    #brings you to document screen
    docss()
def manualsss():
    #this function is page 2 of manuals
    #hides previous page of manual
    packages1.hide(),packages2.hide(), packages3.hide(),packages4.hide(),mannext.hide(),man0.hide(),man1.hide(),man2.hide(),packages.hide()
    #shows new page of manual
    man3.show(),man4.show(),man5.show(),man6.show(),man7.show(), man8.show(),man9.show(),man10.show(),man11.show(),mans2.show()
def manualss():
    #hides previous screens
    settings_button.hide(), doc_button.hide(), address_button.hide(), doc_button.hide(), profile_button.hide()
    #hides previous screen
    doc_button.hide(), manual_button.hide(), hide_all()
    # shows first page of manuals
    packages1.show(),packages2.show(), packages.show(), manual_message.show()
    # shows first page of manuals
    man0.show(),man1.show(),man2.show(),mannext.show(), mans1.show(),packages3.show(), packages4.show()
    # shows first page of manuals
    settings_button1.show(), manual_button1.show(), address_button1.show(), doc_button1.show(), profile_button1.show()
def settingss():
    #hides previous screen
    hide_all()
    #shows settings widgets
    settings_message.show(),  whole_security_message.show(), show_all()
    # shows settings widgets
    settings_button.show(), manual_button.show(), address_button.show(), doc_button.show(), profile_button.show()
    # shows settings widgets
    video_pic_button.show(), video_pic_message.show(), whole_security_button.show(), space_set.show()
def sign_in():
    #iterator
    i = 0
    #ending space for iterator
    length = len(user_list_ary)
    #gets username entered
    a = username_input_box.value + "\n"
    #gets password entered
    b = password_input_box. value + "\n"
    #will look for username password combination
    while i < length :
        #if user name entered is not the current username looked at in list
        if user_list_ary[i] != a :
            #iterate to next element
            i = i + 1
        else:
            #right username, wrong password
            if pass_list_ary[i] != b :
                #iterate
                i = length + 1
            else :
                #will break whileloop
                i = length + 1
                #log in complete opens to settings
                settingss()
def repeated_user():
    #puts all usernames in vector to see if username is already used
    a = user_input_box.value + "\n"
    user_list_ary = []
    with open('user.txt') as f:
        for line in f:
          user_list_ary.append(line)

        #iterator bounds
        length = len(user_list_ary)
        i = 0
        #if it is already used program returns false and repeated_user() terminates
        while i <length:
            if user_list_ary[i] != a:
                pass
            else:
                return False
            i = i + 1

    #invalid user names contain stat, user, vid, pass
    if "stat" in a:
        print("STAT ERROR")
        return False
    if "vid" in a:
        return False
    if "pass" in a:
        return False
    if "user" in a:
        return False
    f.close()
    return True
def trip_zero():

    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 1].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_one():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 2].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_two():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 3].strip())
    print(doc_location)
    os.startfile(doc_location)
    ccap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_three():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 4].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_four():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 5].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_five():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 6].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_six():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 7].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
def trip_seven():
    abc = username_input_box.value + ".txt"
    user_info_ary = []
    with open(abc) as f:
        for line in f:
            user_info_ary.append(line)
    f.close()
    length = len(user_info_ary)
    doc_location = r"C:\Users\Ryan\PycharmProjects\RobUI&"
    doc_location = doc_location.replace('&', r'\'')
    doc_location = doc_location.replace("'", user_info_ary[length - 8].strip())
    print(doc_location)
    os.startfile(doc_location)
    cap = cv2.VideoCapture('Wildlife.wmv')
    while (cap.isOpened()):
        ret, frame = cap.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        cv2.imshow('frame', gray)
        if cv2.waitKey(15) & 0xFF == ord('q'):
            break
    cap.release()
    cv2.destroyAllWindows()
#creates app
app = App(title="Ground Rec Robot", layout="grid", width=530,height=675, bg= "gray50")                                               #creates Ground Rec Robot app
#creates icon buttons
settings_button = PushButton(app, image="settings.png",grid=[0,100], align= "left", command=settingss)
#creates icon buttons
manual_button = PushButton(app, image="manuals.png",grid=[1,100], align="left", command=manualss)
#creates icon buttons
address_button = PushButton(app, image="start.png",grid=[2,100], align="left", command= addressess)
#creates icon buttons
doc_button = PushButton(app, image="doc.png",grid=[3,100], align="left", command=docss)
#creates icon buttons
profile_button = PushButton(app, image="exit.png",grid=[4,100], align="left", command= profilee)
#matches button to app's theme
settings_button.bg="slategray4"
#matches button to app's theme
manual_button.bg="slategray4"
#matches button to app's theme
address_button.bg="slategray4"
#matches button to app's theme
doc_button.bg="slategray4"
#matches button to app's theme
profile_button.bg="slategray4"
#creates vector for usernames
user_list_ary = []
#opens username file
with open('user.txt') as f:
    #begins for loop for username to vector
    for line in f:
        #adds user name to vector
        user_list_ary.append(line)
#close username document
f.close()
#creates password array
pass_list_ary = []
#open file containing passwords
with open('pass.txt') as f:
    #starts a for loop for all the passwords
    for line in f:
        #adds passwords to vector
        pass_list_ary.append(line)
#closes document containing text
f.close()
#Message to alert user they are on create user page
create_user_message = Text(app, text="Create User",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#space for first name to be entered
firstname_message=Text(app, text="First Name", grid=[0,2], color= "black", align="left")
#alert user to type last name
lastname_message=Text(app, text="Last Name", grid=[0,4], color= "black", align= "left")
#alert user to type in user name
user_message_create=Text(app, text="User Name", grid=[0,8], color= "black", align= "left")
#alert user to type password in
password_message_create=Text(app, text="Password", grid=[0,16], color= "black", align= "left")
#space for first name
firstname_input_box = TextBox(app, grid=[1,2], width= 25)
#space for last name
lastname_input_box = TextBox(app, grid=[1,4], width= 25)
#space for username
user_input_box = TextBox(app, grid=[1,8], width= 25)
#space for passeord
pass_input_box = TextBox(app, grid=[1,16], width= 25)
#makes text widgets invisible
firstname_input_box.hide(),lastname_input_box.hide(),user_message_create.hide()
#alerts user to insert qground control location in computer
qground_location_message=Text(app, text="Qground", grid=[0,32], color= "black", align= "left")
#hides text widgets
password_message_create.hide(), qground_location_message.hide()
#creates input box for qground controls location in computer
qground_input_box = TextBox(app, grid=[1,32], width= 25)
#makes widgets invisible
user_input_box.hide(),pass_input_box.hide(),qground_input_box.hide()
#makaes widgets invisible
create_user_message.hide(), firstname_message.hide(), lastname_message.hide()
#finishes creating user
createuser_finished = PushButton(app, text="Create User",grid=[1,128], width=14, command=adduser)
#hides created user button
createuser_finished.hide()
#invalid users are usernames with "stat", "vid" or a repeated user
create_user_error = Text(app, text="Invalid user name",grid=[1,256], width=14, color= "black")
#hides user name error messgae
create_user_error.hide()
#sign in message
signin_message = Text(app, text="Sign In",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#matches background color to apps theme
signin_message.bg= "slategray4"
#indicates user to enter username to sign in
username_message = Text(app, text="Username: ",grid=[0,2], size=16,width=14, color= "black")
#allows user to input username
username_input_box = TextBox(app, grid=[1,2,100,2], width= 25)
#alerts password messgae
password_message = Text(app, text="Password: ",grid=[0,4], size=16,width=14, color= "black")
#allows user to enter signin password
password_input_box = TextBox(app, grid=[1,4,100,4], width= 25)
#creates horizontal spacing
spacing_signin1 = Text(app, text="           ",grid=[0,5,100,10], size=16,width=14, color= "black")
#checks to see if user is a real user
signin_button = PushButton(app, text="Sign in",grid=[1,100], width=14, command=sign_in)
#brings user to be able to create user
createuser = PushButton(app, text="Create User",grid=[0,100], width=14, command=create_user)
#settings message to indicate page on
settings_message = Text(app, text="Settings",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#creates button group for video option
video_pic_button = ButtonGroup(app, options=["Video", "Picture"], grid=[3,2,4,2], horizontal="True", align="left")
#text explaining choices
video_pic_message = Text(app, text="Would you like videos or pictures?", grid= [0,2, 3, 2], align="left")
#Text explaing setting choice
whole_security_message = Text(app, text= "Record whole trip or hazards?", grid=[0,4,3,4], align="left")
#creates button group for video option
whole_security_button = ButtonGroup(app, options=["Whole trip", "Hazards"], grid=[3,4,4,4],horizontal="True", align="left")
#creates spacing for settings icon
space_set = Text(app,text=" ", grid=[0,50],size=10, height=22)
#matches message to app's theme
settings_message.bg="slategray4"
#hides widgets
whole_security_button.hide(), whole_security_message.hide(), video_pic_button.hide(), video_pic_message.hide(), settings_message.hide()
#hides widgets
settings_button.hide(), manual_button.hide(), address_button.hide(), doc_button.hide(), profile_button.hide(),space_set.hide()
#creates text alerting user that they are manual page
manual_message = Text(app, text="Manual",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black",align="left")
#adjust background color to match theme
manual_message.bg="slategray4"
#hides the manual pages message alerting user they are on page
manual_message.hide()
#creates text for the user to be able to read for the manual
packages=Text(app,text="PACKAGES NEEDED", size=12, grid=[0,2,5,2])
#creates text for the user to be able to read for the manual
packages1=Text(app,text="a. RobUI folder", size=10, grid=[0,4,5,4], align="left",height=1)
#creates text for the user to be able to read for the manual
packages2=Text(app,text="b. python/ python IDE", size=10, grid=[0,8,5,8], align="left",height=1)
#creates text for the user to be able to read for the manual
packages3=Text(app,text="c. QGround Control", size=10, grid=[0,16,5,16], align="left",height=1)
#creates text for the user to be able to read for the manual
packages4=Text(app,text="CREATE AN ACCOUNT", size=12, grid=[0,32,5,32])
#creates text for the user to be able to read for the manual
man0=Text(app,text="a. Click create user", size=10, grid=[0,64,5,64], align="left",height=1)
#creates text for the user to be able to read for the manual
man1=Text(app,text="b. Fill in user information", size=10, grid=[0,128,5,128], align="left",height=1)
#creates text for the user to be able to read for the manual
man2=Text(app,text="c. Click create user", size=10, grid=[0,256,5,256], align="left",height=1)
#creates a button for the user to be able to read for the manual
mannext=PushButton(app,text="Next",grid=[4,512],command=manualsss)
#creates text for the user to be able to read for the manual
mans1=Text(app,text=" ", size=9, height=13,grid=[0,1024])
#creates text for the user to be able to read for the manual
man3=Text(app,text="USE ROBOT", size=12, grid=[0,2,5,2])
#creates text for the user to be able to read for the manual
man4=Text(app,text="a. Click on map button", size=10, grid=[0,4,5,4], align="left",height=1)
#creates text for the user to be able to read for the manual
man5=Text(app,text="b. Select locations on Q ground control", size=10, grid=[0,8,5,8], align="left",height=1)
#creates text for the user to be able to read for the manual
man6=Text(app,text="c. Click Start trip", size=10, grid=[0,16,5,16], align="left",height=1)
#creates text for the user to be able to read for the manual
man7=Text(app,text="d. Click Finish trip once complete", size=10, grid=[0,32,5,32], align="left",height=1)
#creates text for the user to be able to read for the manual
man8=Text(app,text="READ DOCUMENTS", size=12, grid=[0,64,5,64])
#creates text for the user to be able to read for the manual
man9=Text(app,text="a. Click on document button", size=10, grid=[0,128,5,128], align="left",height=1)
#creates text for the user to be able to read for the manual
man10=Text(app,text="b. Select desired document", size=10, grid=[0,256,5,256], align="left",height=1)
#creates text for the user to be able to read for the manual
man11=Text(app,text="c. Read document", size=10, grid=[0,512,5,512], align="left",height=1)
#creates text for the user to be able to read for the manual
mans2=Text(app,text=" ",size=10,height=13,grid=[9,1024])
#hides manual widgets
packages.hide(), packages1.hide(),packages2.hide(),packages3.hide(),packages4.hide()
#hides manual widgets
man3.hide(),man4.hide(),man5.hide(),man6.hide(),man7.hide(), man0.hide(),man1.hide(),man2.hide()
#hides manual widgets
mannext.hide(),man8.hide(),man9.hide(),man10.hide(),man11.hide(),mans1.hide(),mans2.hide()
#creates address message to alert address button has been selected
address_message = Text(app, text="Addresses",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#creates button to click to begin trip
start_trip = PushButton(app, text="Start Trip", grid=[0,2,100,2],width=63, command=addressestwo, align="left")
#loads picture of teams logo during travel
logo_picture = Picture(app, image="logo.png", grid=[0,10,100,90], align="left")
#changes height of logo to fill screen
logo_picture.height = 432
#creates button that will stop trip
finish_trip = PushButton(app, text="Finished Trip", grid=[0,2,100,2],width=63, align="left", command=trips)
#changes background color to match app's theme
start_trip.bg="green"
#changes background color to match app's theme
start_trip.text_color="white"
#changes background color to match app's theme
finish_trip.bg = "red"
#changes background color to match app's theme
finish_trip.text_color="white"
#changes background color to match app's theme
address_message.bg="slategray4"
#makes address text/buttons invisible
finish_trip.hide(), address_message.hide(), start_trip.hide(), logo_picture.hide()
#creates text that will let user know they are on document page
documents_message = Text(app, text="Documents",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#creates button that will show  most recent document and video
trip0 = PushButton(app, text="a", grid=[0,2,100,4], align="left", command=trip_zero, width=64)
#creates button that will show 2nd most recent document and video
trip1 = PushButton(app, text="b", grid=[0,6,100,8], align="left", command=trip_one, width=64)
#creates button that will show 3rd most recent document and video
trip2 = PushButton(app, text="c", grid=[0,14,100,16], align="left", command=trip_two, width=64)
#creates button that will show 4th most recent document and video
trip3 = PushButton(app, text="d", grid=[0,30,100,32], align="left", command=trip_three, width=64)
#creates button that will show 5th most recent document and video
trip4 = PushButton(app, text="e", grid=[0,62,100,64], align="left", command=trip_four, width=64)
#creates button that will show 6th most recent document and video
trip5 = PushButton(app, text="f", grid=[0,126,100,128], align="left", command=trip_five, width=64)
#creates button that will show 7th most recent document and video
trip6 = PushButton(app, text="g", grid=[0,254,100,256], align="left", command=trip_six, width=64)
#creates button that will show 8th most recent document and video
trip7 = PushButton(app, text="h", grid=[0,510,100,512], align="left", command=trip_seven, width=64)
#creates button that will bring user to settings page once pushed
settings_button1 = PushButton(app, image="settings.png",grid=[0,2050], align= "left", command=settingss)
#creates button that will bring user to manuals page once pushed
manual_button1 = PushButton(app, image="manuals.png",grid=[1,2050], align="left", command=manualss)
#creates button that will bring user to address page once pushed
address_button1 = PushButton(app, image="start.png",grid=[2,2050], align="left", command= addressess)
#creates button that will bring user to doc page once pushed
doc_button1 = PushButton(app, image="doc.png",grid=[3,2050], align="left", command=docss)
#creates button that will bring user to profile page once pushed
profile_button1 = PushButton(app, image="exit.png",grid=[4,2050], align="left", command=profilee)
#changes color of button to desired color
settings_button1.bg="slategray4"
#changes color of button to desired color
manual_button1.bg="slategray4"
#changes color of button to desired color
address_button1.bg="slategray4"
#changes color of button to desired color
doc_button1.bg="slategray4"
#changes color of button to desired color
profile_button1.bg="slategray4"
#changes background color of the document text to match app's bg
documents_message.bg="slategray4"
#hides trips underneath documents
trip0.hide(),trip1.hide(),trip2.hide(),trip3.hide(),trip4.hide(),trip5.hide(),trip6.hide(),trip7.hide()
#hides second set of icon buttons
settings_button1.hide(),manual_button1.hide(),address_button1.hide(),doc_button1.hide(),profile_button1.hide(),documents_message.hide()
#creates text to serve as spacing underneath documents
trip_space0 = Text(app, text= " ", size = 1, height=205, grid=[0,2000])
#makes the spacing text for documents invisible
trip_space0.hide()
#longitude max of map in degrees
x1 = -96.3070
#longitude min of map in degrees
x2 = -96.4000
#latitude min of map in degrees
y1 = 30.59
#latitude max of map in degrees
y2 = 30.6311
#range oof longitude for the map
dx= x1-x2
#range of latitudes for the map
dy=y2-y1
#vector that stores longitudes of wavepoints that was adesired to be taken
#currently filled with fake data to validate code
long_first_data = (-96.35, -96.35, -96.39)
#vector that stores latitudes of wavepoints that was desired to be taken
#currently filled with fake data to validate code
lat_first_data = (30.6, 30.62, 30.63)
#vector that stores longitude of wavepoints that was actually taken
#currently filled with fake data to validate code
long_end_data = (-96.35,-96.35, -96.37,-96.37, -96.39)
#vector that stores latitudes of wavepoints that was actually taken
#currently filled with fake data to validate code
lat_end_data = (30.6,30.62, 30.625,30.63, 30.63)
#declares, defines and initialize global variable
long_plot_data = []
#declares, defines and initialize global variable
lat_plot_data  = []
#declares, defines and initialize global variable
time_begin_year = 0
#declares, defines and initialize global variable
time_begin_month = 0
#declares, defines and initialize global variable
time_begin_day = 0
#declares, defines and initialize global variable
time_begin_hour = 0
#declares, defines and initialize global variable
time_begin_minute = 0
#declares, defines and initialize global variable
time_begin_second = 0
#declares, defines and initialize global variable
time_beg = 0
#declares, defines and initialize global variable
date_beg = 0
#declares, defines and initialize global variable
time_end_year = 0
#declares, defines and initialize global variable
time_end_month = 0
#declares, defines and initialize global variable
time_end_day = 0
#declares, defines and initialize global variable
time_end_hour = 0
#declares, defines and initialize global variable
time_end_minute = 0
#declares, defines and initialize global variable
time_end_second = 0
#declares, defines and initialize global variable
time_end = 0
#declares, defines and initialize global variable
date_end = 0
#declares, defines and initialize global variable
time_difference = 0
#declares, defines and initialize global variable
distance = 0
#declares, defines and initialize global variable
doc_date = 0
#declares, defines and initialize global variable
doc_name = 0
#creates message at top of profile page to alert user they clicked profile button
profile_message = Text(app, text="Profile",grid=[0,0,100,1], height= 1, size=40,width=14, color= "black")
#makes the top profile message invisible
profile_message.hide()
#changes background color to match app background color
profile_message.bg = "slategray4"
#creates text that will be followed by a stat under profile
trip_total_message = Text(app, text="Trip total", grid=[0,2,3,2], size=10, align="left")
#creates text that will be followed by a stat under profile
time_total_message = Text(app,text= "Time spent on travel", grid=[0,4,3,4], size=10, align= "left")
#creates text that will be followed by a stat under profile
distance_total_message = Text(app,text= "Distance on travel", grid=[0,8,3,8], size=10, align= "left")
#creates text that will be followed by a stat under profile
speed_total_message = Text(app,text= "Average Speed on travel", grid=[0,16,3,16], size=10, align= "left")
#creates text that will be followed by a stat under profile
threats_total_message = Text(app,text= "Potential hazards on travel", grid=[0,32,3,32], size=10, align= "left")
#creates button to logout of account
logout_button = PushButton(app, text="Logout", grid=[4,128], width= 10, command= log_out)
#create space bewteen text and bottom icons
space_profile = Text(app, text= "", height=20, size=1, grid=[0,64])
#hides profile text and buttons until profile is called
time_total_message.hide(), distance_total_message.hide(), speed_total_message.hide(), threats_total_message.hide(), logout_button.hide(), space_profile.hide()
#hides profile text and buttons until profile is called
trip_total_message.hide(),
#boolean to see if profile has been viewed
profile_view = 0
#displays the app, anything under will not be displayed
app.display()